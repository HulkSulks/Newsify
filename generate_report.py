import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

def create_report():
    doc = Document()

    # Define standard styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_heading(text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
        h = doc.add_heading(text, level=level)
        h.alignment = align
        return h

    def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
        p = doc.add_paragraph()
        p.alignment = align
        runner = p.add_run(text)
        runner.bold = bold
        return p

    def add_image(image_path, width_inches=None):
        if os.path.exists(image_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            runner = p.add_run()
            if width_inches:
                runner.add_picture(image_path, width=Inches(width_inches))
            else:
                runner.add_picture(image_path)

    def add_code_block(text):
        table = doc.add_table(rows=1, cols=1, style='Table Grid')
        p = table.rows[0].cells[0].paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        doc.add_paragraph()

    # ================== PAGE 1 ==================
    add_paragraph("\n\n")
    add_paragraph("PROJECT REPORT", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(16)
    add_paragraph("On", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph("HEALTHBRIDGE INDIA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(18)
    add_paragraph("(Government Health Scheme Discovery & Eligibility Platform)", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph("\nSubmitted in", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph("PYTHON PROGRAMMING LAB", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(14)
    add_paragraph("In", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph("Mechatronics Engineering (MCT)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(14)
    add_paragraph("\nBy", align=WD_ALIGN_PARAGRAPH.CENTER)
    
    students = [
        "M VISHNU GANESH (25261A1459)",
        "MANASVI MUGADA (25261A1440)",
        "ABHISHEK VARMA (25261A1433)",
        "YASH DUBEY (25261A1462)",
        "JOHN DANIEL PRATHIK (25261A1449)"
    ]
    for s in students:
        add_paragraph(s, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        
    add_paragraph("\nUnder the Guidance of Subject Faculty of", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph("B.TULASIDAS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph("Assistant Professor, Department of IT\n", align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_image('assets/mgit_logo.png', width_inches=2.0)
    add_paragraph("\n")
    
    add_paragraph("DEPARTMENT OF MECHANICAL AND MECHATRONICS ENGINEERING", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(12)
    add_paragraph("MAHATMA GANDHI INSTITUTE OF TECHNOLOGY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(14)
    add_paragraph("(AUTONOMOUS)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph("(Affiliated to JNTUH, Hyderabad; Eight UG Programs Accredited by NBA;\nAccredited by NAAC with 'A++' Grade)", align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(10)
    
    doc.add_page_break()

    # ================== PAGE 2 ==================
    add_paragraph("Kokapet (Village), Gandipet, Chaitanya Bharathi (P.O.), Rangareddy Dist,\nHYDERABAD — 500075, TELANGANA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph("MAHATMA GANDHI INSTITUTE OF TECHNOLOGY (AUTONOMOUS)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(14)
    add_paragraph("\nCERTIFICATE\n", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(16)
    
    cert_text = (
        "This is to certify that the laboratory project “HEALTHBRIDGE INDIA” has been submitted by "
        "M VISHNU GANESH (25261A1459), MANASVI MUGADA (25261A1440), ABHISHEK VARMA (25261A1433), "
        "YASH DUBEY (25261A1462), JOHN DANIEL PRATHIK (25261A1449) from "
        "Department of Mechanical and Mechatronics Engineering (MCT) to the Department of "
        "Mechanical and Mechatronics Engineering, Mahatma Gandhi Institute of Technology for "
        "B.Tech II-Sem, PYTHON PROGRAMMING LAB course during the academic year 2025–2026."
    )
    p = add_paragraph(cert_text)
    p.paragraph_format.line_spacing = 1.5
    
    add_paragraph("\n\n\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runner = p.add_run("Project Guide:\nMr. TULASI DASU\nAssistant Professor\nDept. of IT")
    runner.bold = True
    
    doc.add_page_break()

    # ================== PAGE 3 ==================
    add_paragraph("DECLARATION\n", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(16)
    
    decl_text = (
        "The successful completion of this project on “HEALTHBRIDGE INDIA” in the Python "
        "Programming Laboratory would not have been possible without the guidance and support of "
        "many individuals. We would like to express our sincere gratitude to our Subject Faculty, Mr. "
        "TULASI DASU, Assistant Professor, Dept. of IT, for his continuous support, patience and "
        "expertise throughout the semester. We also thank the Department of Mechanical and "
        "Mechatronics Engineering and Mahatma Gandhi Institute of Technology for providing the "
        "resources and environment that made this work possible."
    )
    p = add_paragraph(decl_text)
    p.paragraph_format.line_spacing = 1.5
    
    add_paragraph("\n\n\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runner = p.add_run("\n".join(students))
    
    doc.add_page_break()

    # ================== PAGE 4 ==================
    add_paragraph("TABLE OF CONTENTS\n", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(16)
    
    toc = [
        "1. Abstract",
        "2. Introduction",
        "   2.1 Background",
        "   2.2 Problem Statement",
        "   2.3 Objectives",
        "   2.4 Scope",
        "3. System Requirements",
        "   3.1 Software Requirements",
        "   3.2 Hardware Requirements",
        "   3.3 Key Python Dependencies",
        "4. Theory & Technologies Used",
        "   4.1 Tkinter Framework",
        "   4.2 TF-IDF Vectorization",
        "   4.3 Logistic Regression",
        "   4.4 Optical Character Recognition (OCR)",
        "5. System Architecture",
        "6. Application Walkthrough & User Interface",
        "   6.1 The Main Dashboard",
        "   6.2 Detailed Button Functionality",
        "7. Algorithm",
        "8. Flowchart",
        "9. Approach & Methodology",
        "10. Program (Implementation)",
        "11. Output & Results",
        "12. Testing",
        "13. Future Scope & Conclusion",
        "14. References"
    ]
    for item in toc:
        add_paragraph(item).paragraph_format.line_spacing = 1.5
        
    doc.add_page_break()

    # ================== CONTENT PAGES ==================
    # PAGE 5
    add_heading("1. Abstract")
    doc.add_paragraph("Newsify is a desktop GUI application built to identify and flag fake news using traditional Machine Learning techniques. In today's digital era, misinformation spreads rapidly across social networks and news portals, heavily influencing public opinion and, in some cases, inciting panic. Newsify solves this problem by allowing users to instantly verify news articles either by providing a URL or uploading an image of a news snippet.")
    doc.add_paragraph("The application operates by automatically extracting text, cleaning it, and feeding it into a highly trained Logistic Regression model utilizing TF-IDF vectorization. The model was trained on a dataset of over 40,000 real and fake news articles to learn distinct language patterns. It returns a probability score, a confidence level, and a simple template-based explanation for the verdict.")
    doc.add_paragraph("This project demonstrates a clean, modular Python architecture leveraging libraries like Tkinter, scikit-learn, BeautifulSoup, and Pytesseract. It brings AI concepts to a beginner-friendly desktop interface, effectively bridging web scraping, OCR, and Machine Learning.")
    doc.add_page_break()
    
    # PAGE 6
    add_heading("2. Introduction")
    add_heading("2.1 Background", level=2)
    doc.add_paragraph("Fake news refers to misinformation or disinformation presented as legitimate journalism. It influences public perception, elections, and public health. Automated detection tools can assist users in critically evaluating the information they consume. By analyzing patterns in text, machine learning algorithms can classify news credibility much faster than human fact-checkers.")
    add_heading("2.2 Problem Statement", level=2)
    doc.add_paragraph("People often read articles or view screenshots of news on social media without knowing their authenticity. There is a need for a simple, accessible tool that quickly cross-examines the text against patterns of known fake and real news to provide a reliability score. Existing tools are often web-based and require manual text entry. A desktop application capable of processing raw URLs and Images addresses this friction.")
    add_heading("2.3 Objectives", level=2)
    doc.add_paragraph("1. To allow users to verify news seamlessly via a web URL.\n2. To allow users to extract and verify text from an image (screenshot or photo).\n3. To use TF-IDF and Logistic Regression for highly accurate classification.\n4. To provide an intuitive Desktop GUI using Tkinter, featuring Dark Mode.\n5. To maintain a local history of analyzed articles.")
    add_heading("2.4 Scope", level=2)
    doc.add_paragraph("The scope of Newsify covers English-language news detection. It includes the entire pipeline from data ingestion, cleaning, training, to the final production application. It incorporates scraping logic for HTML text extraction and Tesseract OCR for image text extraction, all running locally without reliance on external paid API services.")
    doc.add_page_break()

    # PAGE 7
    add_heading("3. System Requirements")
    add_heading("3.1 Software Requirements", level=2)
    doc.add_paragraph("• Operating System: Windows 10/11, Linux, or macOS\n• Programming Language: Python 3.8 or higher\n• Third-Party Tool: Tesseract OCR Engine (required for image text extraction)\n• IDE: Visual Studio Code or PyCharm (for development)")
    add_heading("3.2 Hardware Requirements", level=2)
    doc.add_paragraph("• Processor: Dual-core 2.0 GHz or higher (Quad-core recommended for faster OCR)\n• RAM: 4 GB minimum (8 GB recommended for loading large models into memory)\n• Storage: ~500 MB free for project, virtual environment, and model pickle files")
    add_heading("3.3 Key Python Dependencies", level=2)
    doc.add_paragraph("The application heavily relies on modern Python libraries installed via pip:\n"
                      "1. scikit-learn: Core machine learning library used for TfidfVectorizer and LogisticRegression.\n"
                      "2. pandas & numpy: Used for large scale data manipulation and CSV handling during model training.\n"
                      "3. joblib: Used for serializing (pickling) the trained model so it doesn't need retraining on every boot.\n"
                      "4. requests & beautifulsoup4: Used to fetch HTML from URLs and parse it into readable text.\n"
                      "5. pytesseract & Pillow: Used for Optical Character Recognition processing on images.\n"
                      "6. nltk: Natural Language Toolkit used for removing English 'stop words'.")
    doc.add_page_break()

    # PAGES 8-9
    add_heading("4. Theory & Technologies Used")
    
    add_heading("4.1 Tkinter Framework", level=2)
    doc.add_paragraph("Tkinter is Python's de-facto standard GUI (Graphical User Interface) package. It is a thin object-oriented layer on top of Tcl/Tk. In Newsify, Tkinter is used to create a responsive, dark-mode lock GUI. The UI implements python's `threading` library, which ensures that while the application is making heavy network requests or running OCR, the Tkinter mainloop does not freeze, allowing the progress bar to animate smoothly.")
    
    add_heading("4.2 TF-IDF Vectorization", level=2)
    doc.add_paragraph("Machine learning algorithms cannot process text natively; they require numerical input. TF-IDF stands for Term Frequency-Inverse Document Frequency. \n\n"
                      "- Term Frequency (TF): Measures how frequently a word appears in a specific document.\n"
                      "- Inverse Document Frequency (IDF): Measures how important a word is. It weighs down highly frequent words across all documents (like 'the', 'is') and scales up rare, unique words.\n\n"
                      "By combining these, Newsify converts an entire news article into a mathematical array representing the unique vocabulary fingerprint of that article.")

    add_heading("4.3 Logistic Regression", level=2)
    doc.add_paragraph("Despite its name, Logistic Regression is a classification algorithm. It is used to predict a binary outcome (0 or 1, Fake or Real). It draws a decision boundary (hyperplane) through the high-dimensional space created by the TF-IDF vectorizer. When a new article is analyzed, the model calculates the mathematical distance from this boundary and uses the sigmoid function to map that distance into a probability score between 0.0 and 1.0.")
    
    add_heading("4.4 Optical Character Recognition (OCR)", level=2)
    doc.add_paragraph("Optical Character Recognition translates images of text into machine-encoded text. Newsify utilizes Google's Tesseract-OCR engine wrapped by the Python library `pytesseract`. When an image is uploaded, it is converted to grayscale using Pillow to enhance contrast. Tesseract then scans the pixels to recognize character shapes, returning a string that can be passed into the ML pipeline.")
    doc.add_page_break()

    # PAGE 10
    add_heading("5. System Architecture")
    doc.add_paragraph("Newsify is architected using a highly modular file structure to separate concerns and ensure maintainability. The codebase is broken down into the following modules:")
    
    arch_table = doc.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = arch_table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Module').bold = True
    hdr_cells[1].paragraphs[0].add_run('Responsibility').bold = True
    
    modules = [
        ("main.py", "Entry Point: Orchestrates the application startup. Checks for trained models and launches the GUI."),
        ("gui.py", "Presentation Layer: Contains FakeNewsDetectorApp. Manages Tkinter widgets, thread dispatching, and UI state."),
        ("detector.py", "Business Logic Layer: Loads ML models. Exposes predict_news() which transforms text and runs Logistic Regression."),
        ("utils.py", "Utility Layer: Pure functions handling external data (BeautifulSoup for HTML, Pytesseract for OCR)."),
        ("history.py", "Data Layer: Manages the in-memory state of past predictions and handles CSV file exporting."),
        ("trainer.py", "Training Script: Standalone script that loads datasets, cleans data, trains the model, and serializes to disk.")
    ]
    
    for mod, resp in modules:
        row_cells = arch_table.add_row().cells
        row_cells[0].text = mod
        row_cells[1].text = resp
    doc.add_page_break()

    # PAGES 11-12
    add_heading("6. Application Walkthrough & User Interface")
    
    add_heading("6.1 The Main Dashboard", level=2)
    doc.add_paragraph("The main dashboard of Newsify is intentionally minimalist. Upon launching the application, the user is greeted by a dark-mode styled interface. The top section contains the application branding. The middle section contains the input entry field, the primary action buttons, and a progress bar. The bottom section is a large, read-only scrolled text area where detailed analysis results are printed.")
    
    add_image('assets/screenshot1.png', width_inches=5.0)
    add_paragraph("Figure 1: The Main Dashboard Interface", align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_heading("6.2 Detailed Button Functionality", level=2)
    
    add_paragraph("Analyze URL Button", bold=True)
    doc.add_paragraph("This button acts upon the text entered in the 'Enter News Article URL' box. When clicked, it disables the UI to prevent double-clicks, starts the indeterminate progress bar, and launches a background thread. The thread sends an HTTP GET request to the URL using a standard User-Agent header. It parses the HTML, finds all <p> tags, concatenates them, cleans the text, and feeds it to the AI. The result is then formatted and displayed. If the URL is invalid or blocked, a graceful error popup is shown.")
    
    add_paragraph("Upload Image Button", bold=True)
    doc.add_paragraph("Clicking this button bypasses the URL input. It opens a native Windows file dialog allowing the user to select an image (.png, .jpg). The selected image is processed by Pillow (converted to grayscale) and read by pytesseract. The extracted text is then passed to the AI exactly as if it were scraped from the web.")
    
    add_paragraph("Clear Button", bold=True)
    doc.add_paragraph("Instantly resets the application state. It wipes the URL entry box, clears the output text area, and updates the status bar back to 'Ready'.")
    
    add_paragraph("History Button", bold=True)
    doc.add_paragraph("Opens a Toplevel secondary window. This window uses a Tkinter Treeview widget to display a tabular history of every prediction made during the current session. It includes columns for Date/Time, Source (URL/Image), Preview Text, Result, and Confidence. From this window, the user can click 'Export as CSV' to save the session history permanently.")
    
    add_paragraph("Copy Button", bold=True)
    doc.add_paragraph("A quality-of-life feature that takes the entire formatted text inside the output area and copies it to the system clipboard, allowing users to easily share the verification results on social media or in messages.")
    
    add_paragraph("Dark Mode Design", bold=True)
    doc.add_paragraph("The application has been permanently locked to a sleek Dark Mode theme. The color palette uses deep slate blues (#0F172A, #1E293B) for backgrounds, with vibrant blue and emerald green for primary action buttons. This provides a modern, premium feel to the application.")
    doc.add_page_break()

    # PAGE 13
    add_heading("7. Algorithm")
    doc.add_paragraph("1. Ingestion: Receive URL string or Image filepath from the user.")
    doc.add_paragraph("2. Extraction: \n   - If URL: Use `requests.get()` and `BeautifulSoup` to extract raw paragraph text.\n   - If Image: Use `Pillow` to grayscale and `pytesseract.image_to_string()` to extract text.")
    doc.add_paragraph("3. Preprocessing: \n   - Convert text to lowercase.\n   - Apply Regex to remove HTTP links and special characters.\n   - Strip excess whitespace.")
    doc.add_paragraph("4. Vectorization: Pass the cleaned string through `loaded_vectorizer.transform()` to yield a sparse numerical matrix.")
    doc.add_paragraph("5. Classification: Pass the matrix into `loaded_model.predict()` and `predict_proba()`.")
    doc.add_paragraph("6. Output Mapping: Convert 0/1 to 'FAKE'/'REAL'. Map the highest probability to a confidence percentage.")
    doc.add_paragraph("7. Explanation Generation: Construct a readable paragraph explaining the model's confidence based on hardcoded templates and word counts.")
    
    add_heading("8. Flowchart")
    doc.add_paragraph("The overall execution flow from user input to result display:")
    add_paragraph("[START] -> [User Input: URL or Image]")
    add_paragraph("             |")
    add_paragraph("     [Is it URL or Image?]")
    add_paragraph("     /                 \\")
    add_paragraph("[Scrape HTML]      [Run Tesseract OCR]")
    add_paragraph("     \\                 /")
    add_paragraph("     [Clean Text (Regex & Lowercase)]")
    add_paragraph("             |")
    add_paragraph("     [TF-IDF Transformation]")
    add_paragraph("             |")
    add_paragraph("  [Logistic Regression Prediction]")
    add_paragraph("             |")
    add_paragraph(" [Format Label, Probabilities, Explanation]")
    add_paragraph("             |")
    add_paragraph("   [Display on GUI & Save to History]")
    add_paragraph("             |")
    add_paragraph("           [END]")
    doc.add_page_break()

    # PAGE 14
    add_heading("9. Approach & Methodology")
    doc.add_paragraph("The development of Newsify followed a highly iterative, bottom-up software engineering methodology:")
    
    doc.add_paragraph("Phase 1: Data Pipeline & Model Training")
    doc.add_paragraph("The foundation of the project was ensuring high accuracy. We began by building `trainer.py`. We sourced the 'Fake and Real News Dataset' from Kaggle containing over 40,000 labeled articles. We implemented pandas to merge, shuffle, and clean the data. We then tested multiple models including Naive Bayes and Logistic Regression. Logistic Regression emerged victorious with an accuracy exceeding 98%. We then used `joblib` to serialize the trained artifacts.")
    
    doc.add_paragraph("Phase 2: Data Acquisition Utilities")
    doc.add_paragraph("With the model trained, the next step was feeding it real-world data. We developed `utils.py` to act as the bridge. We implemented web scraping using BeautifulSoup to navigate DOM trees and extract text. We also integrated Tesseract to allow offline image analysis.")
    
    doc.add_paragraph("Phase 3: Core Application Engine")
    doc.add_paragraph("We built `detector.py` to wrap the serialized model. This decoupled the machine learning mathematics from the user interface. It added safety checks, error handling, and the dynamic explanation generator.")
    
    doc.add_paragraph("Phase 4: Graphical User Interface")
    doc.add_paragraph("Finally, we constructed `gui.py`. We designed a modern grid layout using Tkinter. A major hurdle was GUI freezing during HTTP requests. This was resolved by implementing Python's `threading` module, allowing the analysis functions to run in background daemon threads while the main UI thread animated the progress bar and remained responsive to the user.")
    doc.add_page_break()

    # PAGES 15-18
    add_heading("10. Program (Implementation)")
    doc.add_paragraph("Below are critical excerpts from the project's source code demonstrating the core logic implementation.")
    
    add_heading("10.1 Vectorization and Training (trainer.py)", level=2)
    code1 = """def vectorize_text(train_texts, test_texts):
    # Create the TF-IDF vectorizer (max 50k words)
    vectorizer = TfidfVectorizer(
        max_features=50000,
        stop_words="english",
        ngram_range=(1, 2)
    )
    # Fit on training data and transform it to numbers
    X_train = vectorizer.fit_transform(train_texts)
    X_test = vectorizer.transform(test_texts)
    return vectorizer, X_train, X_test

def train_and_evaluate(X_train, X_test, y_train, y_test):
    lr_model = LogisticRegression(max_iter=1000, random_state=42)
    lr_model.fit(X_train, y_train)
    lr_predictions = lr_model.predict(X_test)
    lr_accuracy = accuracy_score(y_test, lr_predictions)
    print(f"Accuracy: {lr_accuracy * 100:.2f}%")
    return lr_model"""
    add_code_block(code1)
    
    add_heading("10.2 Prediction Engine (detector.py)", level=2)
    code2 = """def predict_news(text):
    cleaned = clean_text(text)
    
    # Convert text into TF-IDF features
    text_features = loaded_vectorizer.transform([cleaned])

    # Get prediction and probabilities
    prediction = loaded_model.predict(text_features)[0]
    probabilities = loaded_model.predict_proba(text_features)[0]

    fake_prob = probabilities[0]
    real_prob = probabilities[1]
    confidence = max(fake_prob, real_prob)
    label = "REAL" if prediction == 1 else "FAKE"

    return {
        "label": label,
        "real_prob": real_prob,
        "fake_prob": fake_prob,
        "confidence": confidence,
        "explanation": generate_explanation(label, confidence, cleaned)
    }"""
    add_code_block(code2)
    
    add_heading("10.3 Multithreaded GUI Execution (gui.py)", level=2)
    code3 = """def analyze_url(self):
    url = self.url_entry.get().strip()
    if not is_valid_url(url):
        messagebox.showerror("Invalid URL", "Check URL format.")
        return

    self.set_buttons_state("disabled")
    self.progress.start(15)

    # Run analysis in background thread
    thread = threading.Thread(target=self.run_url_analysis, args=(url,))
    thread.daemon = True
    thread.start()

def run_url_analysis(self, url):
    try:
        article_text = fetch_url_text(url)
        result = detector.predict_news(article_text)
        
        # Update GUI from background thread
        self.root.after(0, lambda: self.show_prediction(result, "URL"))
        history.add_entry("URL", article_text, result["label"], result["confidence"])
    except Exception as e:
        self.root.after(0, lambda: messagebox.showerror("Error", str(e)))
    self.root.after(0, self.analysis_complete)"""
    add_code_block(code3)
    doc.add_page_break()

    # PAGE 19
    add_heading("11. Output & Results")
    doc.add_paragraph("The application produces a highly accurate verdict on incoming news articles. The results are formatted in a clean, terminal-like output inside the GUI. Below are screenshots demonstrating the application in a live scenario analyzing a BBC News article.")
    
    add_image('assets/screenshot2.png', width_inches=5.0)
    add_paragraph("Figure 2: The BBC Article Used for Analysis", align=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_paragraph("\nWhen the URL is pasted into the application and analyzed, the model accurately classifies the high-quality journalism as REAL NEWS.")
    
    add_image('assets/screenshot3.png', width_inches=5.0)
    add_paragraph("Figure 3: URL Pasted into Newsify Dashboard", align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_image('assets/screenshot4.png', width_inches=5.0)
    add_paragraph("Figure 4: The Final Prediction Output showing High Confidence", align=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_page_break()

    # PAGE 20
    add_heading("12. Testing")
    doc.add_paragraph("The application was subjected to rigorous manual testing to ensure reliability, graceful error handling, and GUI responsiveness. The representative test cases are documented below:")
    
    table = doc.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Test Case').bold = True
    hdr_cells[1].paragraphs[0].add_run('Input').bold = True
    hdr_cells[2].paragraphs[0].add_run('Expected Result').bold = True
    hdr_cells[3].paragraphs[0].add_run('Status').bold = True
    
    test_cases = [
        ("Valid URL", "A standard BBC News URL", "Progress bar starts, text is fetched, model predicts REAL, history is updated", "Passed"),
        ("Invalid URL", "'htp:/broken-link'", "Regex validator catches format error immediately and shows warning dialog", "Passed"),
        ("404 Error", "Perfectly formatted URL leading to a dead page", "requests.get() throws 404, caught gracefully showing error dialog", "Passed"),
        ("Valid Image OCR", "Clear screenshot of an article", "Pytesseract extracts text accurately, model runs prediction", "Passed"),
        ("Non-Image Upload", "Uploading a .pdf or .txt file instead of an image", "File dialog forces image extensions, preventing selection", "Passed"),
        ("Concurrency", "Clicking 'Analyze URL' multiple times rapidly", "UI disables buttons upon first click, preventing thread stacking", "Passed")
    ]
    
    for tc, inp, exp, status in test_cases:
        row_cells = table.add_row().cells
        row_cells[0].text = tc
        row_cells[1].text = inp
        row_cells[2].text = exp
        row_cells[3].text = status
    doc.add_page_break()

    # PAGE 21
    add_heading("13. Future Scope & Conclusion")
    add_heading("13.1 Future Scope", level=2)
    doc.add_paragraph("1. Deep Learning Integration: Replace the Logistic Regression model with a fine-tuned Transformer model like BERT or RoBERTa. This would allow the application to understand semantic context and sarcasm, rather than just analyzing word frequency.")
    doc.add_paragraph("2. Multi-Language Support: Expand the training dataset and OCR capabilities to support regional languages like Hindi or Telugu, massively expanding the tool's impact in India.")
    doc.add_paragraph("3. Fact-Checking APIs: Integrate APIs like the Google Fact Check Explorer to cross-reference specific claims in real-time against verified internet sources.")
    doc.add_paragraph("4. Browser Extension: Repackage the backend logic into a Chrome Extension that automatically highlights fake news natively in the browser.")
    
    add_heading("13.2 Conclusion", level=2)
    doc.add_paragraph("Newsify successfully demonstrates how Python can be utilized to build a complete, end-to-end Machine Learning desktop application. By combining the data-processing power of scikit-learn, the web scraping capabilities of BeautifulSoup, the OCR magic of Tesseract, and the interface design of Tkinter, this project solves a genuine societal problem. It provides an accessible tool for citizens to evaluate the news they consume, promoting digital literacy. The modular architecture ensures the codebase is highly maintainable and serves as a strong foundation for future AI enhancements.")
    doc.add_page_break()

    # PAGE 22
    add_heading("14. References")
    doc.add_paragraph("• Scikit-Learn Documentation — https://scikit-learn.org/stable/")
    doc.add_paragraph("• Python Tkinter Guide — https://docs.python.org/3/library/tkinter.html")
    doc.add_paragraph("• BeautifulSoup Documentation — https://www.crummy.com/software/BeautifulSoup/bs4/doc/")
    doc.add_paragraph("• Tesseract OCR Engine — https://github.com/tesseract-ocr/tesseract")
    doc.add_paragraph("• PyTesseract Wrapper — https://pypi.org/project/pytesseract/")
    doc.add_paragraph("• Kaggle Fake and Real News Dataset by Clément Bisaillon — https://www.kaggle.com/datasets/clmentbisaillon/fake-and-real-news-dataset")
    doc.add_paragraph("• Pandas Data Manipulation — https://pandas.pydata.org/docs/")

    doc.save('Newsify_Project_Report_v3.docx')

if __name__ == "__main__":
    create_report()
